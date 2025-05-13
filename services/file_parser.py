import tempfile
import os 
import docx
import re
from docx import Document
import fitz # using pymupdf instread of #from PyPDF2 

async def create_temp_file(input_file):
    """ Create a temporary file from uploaded file"""
    try:            
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{input_file.filename.split('.')[-1]}") as temp_file:
            input_file_path = temp_file.name
            resume_bytes = await input_file.read()
            temp_file.write(resume_bytes)
        return input_file_path
    except Exception as e:
        RuntimeError(f"Failed to create temporary file: {str(e)}")

def get_text(inputfile):
    if inputfile is not None: 
        raw_text = get_text_from_file(inputfile) if inputfile is not None else None
        return raw_text
    return ""

def get_text_from_file(uploaded_file):
    """Extract text from a PDF, docx, txt files when a user uploads a file."""
    ret_text =""
    error = None
    try:
        ret_text =""
        error = None
        # For Text files
        if uploaded_file is not None and uploaded_file.name.endswith('.txt'):
            # Read the file as bytes
            file_content = uploaded_file.getvalue()
            # Decode the bytes to string
            ret_text = file_content.decode("utf-8")
            return ret_text , error
        elif uploaded_file is not None and uploaded_file.name.endswith('.docx'):
            # Read the file
            if uploaded_file is not None:
                doc = Document(uploaded_file)
                ret_text = "\n".join([para.text for para in doc.paragraphs])
                return ret_text, error

        elif uploaded_file is not None and uploaded_file.name.endswith('.pdf'):
            if uploaded_file is not None:
                with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
                    ret_text = "\n".join([page.get_text() for page in doc])
                    return ret_text, error
        else:
            error= "Unsupported file type. Please upload a .txt, .docx, or .pdf file."
            return None, error
    except Exception as e:
        error = f"Error while processing file: {e}"
        return None, error

def get_cleantext(text:str):
    """ Clean and Normalize text""" 
    # try:
    #     text = text.encode("utf-8").decode("utf-8")  # Ensure the text is in UTF-8 format
    # except UnicodeDecodeError:
    #     text = text.encode("latin-1").decode("utf-8", errors="ignore")

    text = re.sub(r"[^a-zA-Z0-9\s\-\.\\n\:@$%,]", "", text)  # Removes special characters
    # text = re.sub(r"\\t", " ", text)  # Removes special characters
    # text = " ".join(text.split())  # Removes extra spaces
    # text =  text.lower()  # Converts to lowercase 
    return text

class FileParser():
    """ Provide a file of pdf, dox, txt and a raw text will be returned
    """ 
    def __init__(self): 
        self.raw_text = ""

    def get_raw_text(self, input_file): 
        try:
            if type(input_file) == bytes: 
                self.raw_text = get_text(input_file)
            elif type(input_file) == str: 
                return self.parse_file(input_file) 
        except Exception as e:
            raise ValueError( f"FileParser.get_raw_text: Exception: {e}")
        return  self.raw_text 
        
    def parse_file(self, file_path):
        """Extract text from resume file (PDF or DOCX)"""
        text = ""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                with fitz.open(file_path, filetype="pdf") as doc:
                    text = "\n".join([page.get_text() for page in doc]) 
                # Following for PyPDF2
                # with open(file_path, 'rb') as file:
                #     reader = pyPDF2(file) 
                #     for page_num in range(len(reader.pages)):
                #         text += reader.pages[page_num].extract_text() + "\n"
            elif file_ext in ['.docx', '.doc']:
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            elif file_ext in ['.txt']:
                with open(file_path, "r", encoding="utf-8") as file:
                    return file.read() 
            else:
                raise ValueError(f"FileParser.parse_file: Unsupported file format: {file_ext}")
        except Exception as e:
            raise ValueError( f"FileParser.parse_file: Exception: {e}")
        
        text = get_cleantext(text)
        return text